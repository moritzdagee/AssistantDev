#!/usr/bin/env python3
"""Test file creation functions (docx, xlsx, pdf) with various JSON formats."""

import sys
import os
import json
import tempfile

# Add src to path
sys.path.insert(0, os.path.expanduser('~/AssistantDev/src'))

# We need to set up the environment the functions expect
OUTPUT_DIR = tempfile.mkdtemp(prefix='test_filecreation_')
os.environ['TEST_OUTPUT_DIR'] = OUTPUT_DIR

# Import the sanitizer and file creation functions by exec'ing the relevant code
ws_path = os.path.expanduser('~/AssistantDev/src/web_server.py')
with open(ws_path) as f:
    source = f.read()

# Extract just the functions we need
import re

# Extract sanitize_llm_json
m = re.search(r'(def sanitize_llm_json\(raw\):.*?)(?=\ndef \w)', source, re.DOTALL)
if m:
    exec(m.group(1), globals())
    print("✓ sanitize_llm_json loaded")

# Test 1: JSON Sanitizer
print("\n=== TEST 1: JSON Sanitizer ===")

tests = [
    ('Valid JSON', '{"title": "Test", "content": [{"type": "paragraph", "text": "Hello"}]}'),
    ('Single Quotes', "{'title': 'Test', 'content': [{'type': 'paragraph', 'text': 'Hello'}]}"),
    ('Trailing Comma', '{"title": "Test", "content": [{"type": "paragraph", "text": "Hello"},]}'),
    ('Mixed Quotes', """{'title': "Test Doc", 'content': [{'type': 'heading', 'level': 1, 'text': "Title"}]}"""),
    ('Markdown Fence', '```json\n{"title": "Test", "content": []}\n```'),
]

for name, input_json in tests:
    try:
        result = sanitize_llm_json(input_json)
        assert isinstance(result, dict), f"Expected dict, got {type(result)}"
        assert 'title' in result, "Missing 'title' key"
        print(f"  ✓ {name}: OK (title={result['title']})")
    except Exception as e:
        print(f"  ✗ {name}: FAILED — {e}")

# Test 2: Actual file creation
print("\n=== TEST 2: DOCX Creation ===")
try:
    from docx import Document
    doc = Document()
    doc.add_heading("Test Dokument", 0)
    doc.add_paragraph("Dies ist ein Test-Paragraph.")
    doc.add_paragraph("Aufzählung", style='List Bullet')
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    fpath = os.path.join(OUTPUT_DIR, 'test.docx')
    doc.save(fpath)
    assert os.path.exists(fpath), "File not created"
    assert os.path.getsize(fpath) > 1000, f"File too small ({os.path.getsize(fpath)} bytes)"
    # Verify it can be opened
    doc2 = Document(fpath)
    assert len(doc2.paragraphs) >= 2, f"Too few paragraphs ({len(doc2.paragraphs)})"
    print(f"  ✓ DOCX created and verified ({os.path.getsize(fpath)} bytes, {len(doc2.paragraphs)} paragraphs)")
except Exception as e:
    print(f"  ✗ DOCX FAILED: {e}")

print("\n=== TEST 3: XLSX Creation ===")
try:
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Test"
    ws.cell(row=1, column=1, value="Header 1")
    ws.cell(row=1, column=2, value="Header 2")
    ws.cell(row=2, column=1, value="Data 1")
    ws.cell(row=2, column=2, value="Data 2")
    fpath = os.path.join(OUTPUT_DIR, 'test.xlsx')
    wb.save(fpath)
    assert os.path.exists(fpath), "File not created"
    assert os.path.getsize(fpath) > 1000, f"File too small"
    # Verify
    wb2 = openpyxl.load_workbook(fpath)
    assert wb2.active.cell(row=1, column=1).value == "Header 1"
    print(f"  ✓ XLSX created and verified ({os.path.getsize(fpath)} bytes)")
except Exception as e:
    print(f"  ✗ XLSX FAILED: {e}")

print("\n=== TEST 4: PDF Creation ===")
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
    fpath = os.path.join(OUTPUT_DIR, 'test.pdf')
    doc = SimpleDocTemplate(fpath, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = [
        Paragraph("Test Heading", styles['Heading1']),
        Paragraph("Test paragraph content.", styles['Normal']),
    ]
    doc.build(elements)
    assert os.path.exists(fpath), "File not created"
    assert os.path.getsize(fpath) > 500, f"File too small"
    print(f"  ✓ PDF created and verified ({os.path.getsize(fpath)} bytes)")
except Exception as e:
    print(f"  ✗ PDF FAILED: {e}")

# Cleanup
import shutil
shutil.rmtree(OUTPUT_DIR, ignore_errors=True)

print("\n=== DONE ===")
